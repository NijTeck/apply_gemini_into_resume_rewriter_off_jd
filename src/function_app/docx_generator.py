import io
import re
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("jobapplier-docx-generator")

# --- Helper Functions for Formatting ---

def set_margins(doc, top=0.5, bottom=0.5, left=0.5, right=0.5):
    """Set document margins in inches."""
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)

def set_default_font(doc, font_name="Calibri", font_size=Pt(10)):
    """Set the default document font."""
    style = doc.styles["Normal"]
    font = style.font
    font.name = font_name
    font.size = font_size
    # Set complex script font as well for compatibility
    rpr = style.element.xpath(".//w:rPr")[0]
    r_fonts = rpr.find(qn("w:rFonts"))
    if r_fonts is not None:
        r_fonts.set(qn("w:cs"), font_name)

def add_styled_paragraph(doc, text, style_name=None, size=None, bold=False, italic=False, align=None, space_after=Pt(4), space_before=Pt(0)):
    """Add a paragraph with specific styling."""
    para = doc.add_paragraph()
    if style_name:
        para.style = doc.styles[style_name]
    run = para.add_run(text)
    if size:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    if align:
        para.alignment = align
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    return para

def add_bullet_point(doc, text, level=0, space_after=Pt(2)):
    """Add a bullet point with optional indentation."""
    # Use standard List Bullet style, adjust indentation if needed
    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = Pt(0)
    
    # Split the text into sentences for better formatting of long bullets
    sentences = text.split('. ')
    first_sentence = True
    
    for i, sentence in enumerate(sentences):
        if not sentence.strip():
            continue
            
        # Add period back except for the last sentence if it doesn't end with punctuation
        if i < len(sentences) - 1 or sentence[-1] not in ['.', '!', '?']:
            sentence = sentence + '.'
            
        # First sentence is bold to highlight the key point
        if first_sentence:
            run = para.add_run(sentence + ' ')
            if level == 0:  # Only bold first sentence of main bullets
                run.bold = True
            first_sentence = False
        else:
            para.add_run(sentence + ' ')
    
    # Ensure bullet font matches default
    for run in para.runs:
        run.font.name = doc.styles["Normal"].font.name
        run.font.size = doc.styles["Normal"].font.size
    
    return para

def add_horizontal_line(paragraph):
    """Add a horizontal line below a paragraph."""
    p_borders = OxmlElement("w:pBdr")
    bottom_border = OxmlElement("w:bottom")
    bottom_border.set(qn("w:val"), "single")
    bottom_border.set(qn("w:sz"), "6") # Corresponds to 0.75pt
    bottom_border.set(qn("w:space"), "1")
    bottom_border.set(qn("w:color"), "auto")
    p_borders.append(bottom_border)
    paragraph._p.get_or_add_pPr().append(p_borders)

# --- Main Function ---

def text_to_docx(resume_text: str) -> bytes:
    """
    Convert structured resume text (with markers) to a professionally formatted DOCX file.
    
    Args:
        resume_text: The structured text of the resume with markers like [NAME], [CONTACT], etc.
    
    Returns:
        A bytes object containing the DOCX file.
    """
    try:
        doc = Document()
        set_margins(doc) # Use narrower margins: 0.5 inch all around
        set_default_font(doc, font_name="Calibri", font_size=Pt(10)) # Slightly smaller font for more content
        
        # Customize the built-in styles
        # Make bullet points more compact
        if "List Bullet" in doc.styles:
            bullet_style = doc.styles["List Bullet"]
            bullet_style.paragraph_format.space_after = Pt(2)
            bullet_style.paragraph_format.space_before = Pt(0)
            bullet_style.paragraph_format.line_spacing = 1.0
        
        lines = resume_text.strip().split("\n")
        
        # Define regex patterns for markers
        patterns = {
            "NAME": re.compile(r"^\s*\[NAME\]\s*(.*)", re.IGNORECASE),
            "CONTACT": re.compile(r"^\s*\[CONTACT\]\s*(.*)", re.IGNORECASE),
            "SUMMARY": re.compile(r"^\s*\[SUMMARY\]\s*(.*)", re.IGNORECASE),
            "SECTION_HEADER": re.compile(r"^\s*\[SECTION_HEADER\]\s*(.*)", re.IGNORECASE),
            "JOB_TITLE": re.compile(r"^\s*\[JOB_TITLE\]\s*(.*)", re.IGNORECASE),
            "COMPANY": re.compile(r"^\s*\[COMPANY\]\s*(.*)", re.IGNORECASE),
            "DATES": re.compile(r"^\s*\[DATES\]\s*(.*)", re.IGNORECASE),
            "LOCATION": re.compile(r"^\s*\[LOCATION\]\s*(.*)", re.IGNORECASE),
            "BULLET": re.compile(r"^\s*\[BULLET\]\s*(.*)", re.IGNORECASE),
            "SKILL_CATEGORY": re.compile(r"^\s*\[SKILL_CATEGORY\]\s*(.*)", re.IGNORECASE),
            "SKILLS": re.compile(r"^\s*\[SKILLS\]\s*(.*)", re.IGNORECASE),
            "EDUCATION_DEGREE": re.compile(r"^\s*\[EDUCATION_DEGREE\]\s*(.*)", re.IGNORECASE),
            "EDUCATION_SCHOOL": re.compile(r"^\s*\[EDUCATION_SCHOOL\]\s*(.*)", re.IGNORECASE),
            "EDUCATION_DATES": re.compile(r"^\s*\[EDUCATION_DATES\]\s*(.*)", re.IGNORECASE),
            "EDUCATION_DETAILS": re.compile(r"^\s*\[EDUCATION_DETAILS\]\s*(.*)", re.IGNORECASE),
        }

        last_job_para = None # To align dates/location
        last_edu_para = None # For education entries

        # Count the number of bullet points to ensure we have enough
        bullet_count = 0
        for line in lines:
            if re.match(r"^\s*\[BULLET\]", line, re.IGNORECASE):
                bullet_count += 1
        
        logger.info(f"Resume contains {bullet_count} bullet points")
        if bullet_count < 15:
            logger.warning(f"Resume may not have enough bullet points ({bullet_count}) to fill two pages")

        for line in lines:
            line = line.strip()
            if not line:
                continue # Skip empty lines

            matched = False
            for marker_type, pattern in patterns.items():
                match = pattern.match(line)
                if match:
                    text = match.group(1).strip()
                    if not text:
                        matched = True # Handle empty markers if needed
                        break
                        
                    # Apply formatting based on marker type
                    if marker_type == "NAME":
                        p = add_styled_paragraph(doc, text, size=Pt(16), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(2))
                    elif marker_type == "CONTACT":
                        p = add_styled_paragraph(doc, text, size=Pt(9), align=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(6))
                        add_horizontal_line(p) # Add line after contact info
                    elif marker_type == "SUMMARY":
                        # Handle multi-line summary
                        summary_lines = text.split('\n')
                        for i, summary_line in enumerate(summary_lines):
                            if i == 0 and "US Citizen" in summary_line:
                                # Special handling for US Citizen line
                                p = add_styled_paragraph(doc, summary_line, bold=True, space_after=Pt(2))
                            else:
                                # Regular summary text
                                p = add_styled_paragraph(doc, summary_line, space_after=Pt(2))
                    elif marker_type == "SECTION_HEADER":
                        p = add_styled_paragraph(doc, text.upper(), size=Pt(11), bold=True, space_after=Pt(3), space_before=Pt(6))
                        add_horizontal_line(p)
                    elif marker_type == "JOB_TITLE":
                        last_job_para = add_styled_paragraph(doc, text, bold=True, space_after=Pt(0))
                    elif marker_type == "COMPANY":
                        # If previous was job title, add company to same line, otherwise new para
                        if last_job_para and last_job_para.text.strip():
                            last_job_para.add_run(f" | {text}").italic = True
                        else:
                             last_job_para = add_styled_paragraph(doc, text, italic=True, space_after=Pt(0))
                    elif marker_type == "DATES":
                         # If previous was job title/company, add dates aligned right
                        if last_job_para:
                            tab_stops = last_job_para.paragraph_format.tab_stops
                            tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                            last_job_para.add_run("\t" + text)
                            last_job_para.paragraph_format.space_after = Pt(2) # Reset space after adding date
                            last_job_para = None # Reset for next job block
                        else:
                            add_styled_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(2))
                    elif marker_type == "LOCATION":
                        # Similar logic for location if needed, potentially on a new line or combined
                        add_styled_paragraph(doc, text, italic=True, size=Pt(9), space_after=Pt(2))
                    elif marker_type == "BULLET":
                        # Use the enhanced bullet point function for better formatting of detailed bullets
                        add_bullet_point(doc, text, space_after=Pt(2))
                    elif marker_type == "SKILL_CATEGORY":
                        add_styled_paragraph(doc, text + ": ", bold=True, space_after=Pt(0), space_before=Pt(3))
                    elif marker_type == "SKILLS":
                        # Add skills to the same line as the category if possible
                        if doc.paragraphs[-1].text.endswith(": "):
                            doc.paragraphs[-1].add_run(text)
                            doc.paragraphs[-1].paragraph_format.space_after = Pt(3)
                        else:
                            add_styled_paragraph(doc, text, space_after=Pt(3))
                    elif marker_type == "EDUCATION_DEGREE":
                        last_edu_para = add_styled_paragraph(doc, text, bold=True, space_after=Pt(0))
                    elif marker_type == "EDUCATION_SCHOOL":
                         if last_edu_para and last_edu_para.text.strip():
                            last_edu_para.add_run(f", {text}").italic = True
                         else:
                            last_edu_para = add_styled_paragraph(doc, text, italic=True, space_after=Pt(0))
                    elif marker_type == "EDUCATION_DATES":
                        if last_edu_para and last_edu_para.text.strip():
                            tab_stops = last_edu_para.paragraph_format.tab_stops
                            tab_stops.add_tab_stop(Inches(6.5), alignment=WD_ALIGN_PARAGRAPH.RIGHT)
                            last_edu_para.add_run("\t" + text)
                            last_edu_para.paragraph_format.space_after = Pt(2)
                            last_edu_para = None
                        else:
                            add_styled_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=Pt(2))
                    elif marker_type == "EDUCATION_DETAILS":
                        add_styled_paragraph(doc, text, size=Pt(9), italic=True, space_after=Pt(3))
                        
                    matched = True
                    break # Move to next line once a marker is matched
            
            # If no marker matched, treat as a standard paragraph (e.g., description under job title)
            if not matched and line:
                add_styled_paragraph(doc, line, space_after=Pt(3))
                last_job_para = None # Reset job para tracking if it was description text
                last_edu_para = None # Reset edu para tracking

        # Save the document to a bytes buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        logger.info("Successfully generated formatted DOCX")
        return buffer.getvalue()
    
    except Exception as e:
        logger.error(f"Error generating DOCX: {str(e)}")
        # Consider returning a simple DOCX with the error or raising the exception
        # For now, re-raise to indicate failure
        raise

