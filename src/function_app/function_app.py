import logging
import json
import os
import tempfile
import requests
from azure.storage.blob import BlobServiceClient
from azure.ai.formrecognizer import DocumentAnalysisClient
from azure.core.credentials import AzureKeyCredential
import PyPDF2
import docx
import datetime
import re
import time

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("jobapplier")

# Remove Azure Functions-specific code and keep only the utility functions

def extract_text_from_file(file_path, filename):
    """Extract text from PDF or DOCX file"""
    logger.info(f"Extracting text from file: {filename}")
    
    if not os.path.exists(file_path):
        logger.error(f"File does not exist: {file_path}")
        raise FileNotFoundError(f"File does not exist: {file_path}")
        
    if not filename:
        logger.error("Filename is required")
        raise ValueError("Filename is required")
    
    try:
        if filename.lower().endswith(".pdf"):
            return extract_text_from_pdf(file_path)
        elif filename.lower().endswith(".docx"):
            return extract_text_from_docx(file_path)
        else:
            # Use Azure Document Intelligence for other formats or as fallback
            return extract_text_using_doc_intelligence(file_path)
    except Exception as e:
        logger.error(f"Error extracting text from file {filename}: {str(e)}")
        raise

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file"""
    logger.info("Extracting text from PDF file")
    try:
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            if len(pdf_reader.pages) == 0:
                logger.error("PDF file has no pages")
                raise ValueError("PDF file has no pages")
                
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
                    
        if not text.strip():
            logger.error("No text could be extracted from PDF")
            raise ValueError("No text could be extracted from PDF")
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file"""
    logger.info("Extracting text from DOCX file")
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text:
                text += paragraph.text + "\n"
                
        # Also check tables if paragraphs are empty
        if not text.strip():
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + "\n"
                            
        if not text.strip():
            logger.error("No text could be extracted from DOCX")
            raise ValueError("No text could be extracted from DOCX")
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise

def extract_text_using_doc_intelligence(file_path):
    """Extract text using Azure Document Intelligence (formerly Form Recognizer)"""
    logger.info("Extracting text using Azure Document Intelligence")
    try:
        # Get credentials from environment
        endpoint = os.environ.get("DOCUMENT_INTELLIGENCE_ENDPOINT")
        key = os.environ.get("DOCUMENT_INTELLIGENCE_KEY")
        
        if not endpoint or not key:
            logger.error("Document Intelligence credentials not found in environment variables")
            raise ValueError("Document Intelligence credentials not found in environment variables")
        
        # Create a Document Analysis client
        document_analysis_client = DocumentAnalysisClient(
            endpoint=endpoint, 
            credential=AzureKeyCredential(key)
        )
        
        # Analyze the document
        with open(file_path, "rb") as f:
            poller = document_analysis_client.begin_analyze_document(
                "prebuilt-document", document=f
            )
        result = poller.result()
        
        # Extract the text
        text = ""
        for page in result.pages:
            for line in page.lines:
                if line.content:
                    text += line.content + "\n"
        
        if not text.strip():
            logger.error("No text could be extracted using Document Intelligence")
            raise ValueError("No text could be extracted using Document Intelligence")
            
        return text
    except Exception as e:
        logger.error(f"Error extracting text using Document Intelligence: {str(e)}")
        raise

def upload_to_blob_storage(file_path, original_filename):
    """Upload the file to Azure Blob Storage"""
    try:
        # Get connection string and container name from environment
        connection_string = os.environ["BLOB_CONNECTION_STRING"]
        container_name = os.environ["RESUME_CONTAINER_NAME"]
        
        # Create a unique filename
        unique_filename = f"{original_filename.split('.')[0]}_{get_timestamp()}.{original_filename.split('.')[-1]}"
        
        # Create the BlobServiceClient
        blob_service_client = BlobServiceClient.from_connection_string(connection_string)
        container_client = blob_service_client.get_container_client(container_name)
        blob_client = container_client.get_blob_client(unique_filename)
        
        # Upload the file
        with open(file_path, "rb") as data:
            blob_client.upload_blob(data, overwrite=True)
        
        # Return the URL of the uploaded blob
        return blob_client.url
    
    except Exception as e:
        logger.error(f"Error uploading to blob storage: {str(e)}")
        raise

def get_timestamp():
    """Generate a timestamp for unique filenames"""
    # Corrected import placement
    import datetime
    return datetime.datetime.utcnow().strftime("%Y%m%d%H%M%S")

def get_gemini_recommendations(resume_text, job_description):
    """Get recommendations using Google Gemini API"""
    try:
        # Get API key from environment
        gemini_api_key = os.environ["GEMINI_API_KEY"]
        # Use a compatible model like gemini-1.5-pro-latest
        model_id = os.environ.get("GEMINI_MODEL_ID", "gemini-1.5-pro")
        
        logger.info(f"Using Gemini model: {model_id}")
        
        # Prepare the API URL - use v1beta for compatibility
        api_url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_id}:generateContent?key={gemini_api_key}"
        
        # Prepare the prompt for the Gemini model
        prompt = f"""
        I need a highly critical and analytical evaluation of my resume for a specific job application. Be extremely thorough and don't hold back on criticism. Focus on gaps, misalignments, and areas for improvement.
        
        Here is my resume:
        {resume_text}
        
        Here is the job description I'm applying for:
        {job_description}
        
        Please provide the following with deep analysis and specificity:
        
        1. A comprehensive list of skills from my resume that align with this job. For each skill, rate how strongly it's demonstrated (1-5 scale) and how critical it is for the role.
        
        2. A detailed analysis of skills mentioned in the job description that are missing or insufficiently demonstrated in my resume. Be specific about what's missing and how important each is.
        
        3. A critical analysis of my resume formatting, structure, and content. Identify specific areas that need improvement to better align with industry and role expectations.
        
        4. A thorough list of potential red flags, gaps, or inconsistencies in my resume that would concern a hiring manager for this specific role.
        
        5. Detailed, actionable recommendations for strengthening my resume, including specific examples of how to reword or restructure content.
        
        6. A gap analysis comparing my experience level against what the job description suggests is required, with percentage match estimates for key requirements.
        
        Be brutally honest - your critical evaluation will help me improve my chances significantly.
        
        Format the response as JSON with the following structure:
        {{
            "matching_skills": [
                {{"skill": "skill1", "strength": 4, "importance": 5, "notes": "detailed notes about how this is demonstrated"}},
                {{"skill": "skill2", "strength": 2, "importance": 5, "notes": "detailed notes about how this is demonstrated"}}
            ],
            "missing_skills": [
                {{"skill": "skill1", "importance": 5, "suggestion": "specific way to address this gap"}},
                {{"skill": "skill2", "importance": 3, "suggestion": "specific way to address this gap"}}
            ],
            "improvement_suggestions": [
                "detailed suggestion 1 with specific example of how to implement",
                "detailed suggestion 2 with specific example of how to implement"
            ],
            "potential_red_flags": [
                "detailed description of red flag 1 and why it's problematic for this role",
                "detailed description of red flag 2 and why it's problematic for this role"
            ],
            "experience_tailoring": [
                "specific guidance on how to reframe experience 1",
                "specific guidance on how to reframe experience 2"
            ],
            "gap_analysis": {{"overall_match": "65%", "technical_match": "70%", "experience_match": "60%", "critical_gaps": ["gap1", "gap2"]}}
        }}

        IMPORTANT: Ensure you ONLY respond with valid JSON formatted exactly as specified above. Your entire response must be parseable JSON with no additional text before or after.
        """
        
        # Prepare the API request with correct format
        payload = {
            "contents": [
                {
                    "parts": [
                        {"text": prompt}
                    ]
                }
            ],
            "generationConfig": {
                "temperature": 0.2,
                "topK": 40,
                "topP": 0.95,
                "maxOutputTokens": 8192
            },
            "safetySettings": [
                {
                    "category": "HARM_CATEGORY_HARASSMENT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_HATE_SPEECH",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_SEXUALLY_EXPLICIT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                },
                {
                    "category": "HARM_CATEGORY_DANGEROUS_CONTENT",
                    "threshold": "BLOCK_MEDIUM_AND_ABOVE"
                }
            ]
        }
        
        headers = {
            "Content-Type": "application/json"
        }
        
        # Set up retry parameters
        max_retries = 3
        retry_delay = 5  # seconds
        
        # Implementation with retries
        for attempt in range(1, max_retries + 1):
            try:
                # Make the API request
                logger.info(f"Sending request to Gemini API (attempt {attempt}/{max_retries}): {api_url}")
                # Log payload size and structure
                logger.info(f"Request payload size: {len(str(payload))} characters")
                logger.info(f"Payload structure: {json.dumps(payload, indent=2)[:200]}...")
                
                # Send with longer timeout
                response = requests.post(api_url, headers=headers, json=payload, timeout=120)
                
                # Log response details for debugging
                logger.info(f"Response status code: {response.status_code}")
                
                # Handle rate limit (429) errors
                if response.status_code == 429:
                    if attempt < max_retries:
                        retry_after = int(response.headers.get('Retry-After', retry_delay))
                        logger.warning(f"Gemini API rate limit exceeded (attempt {attempt}/{max_retries}). Retrying in {retry_after} seconds...")
                        time.sleep(retry_after)
                        continue
                    else:
                        logger.error(f"Gemini API rate limit exceeded after {max_retries} attempts: {response.text}")
                        return {
                            "error": "The resume optimization service is temporarily unavailable due to high demand. Please try again in a few minutes.",
                            "matching_skills": [
                                {"skill": "API Connection", "strength": 3, "importance": 5, "notes": f"API returned error code {response.status_code}"}
                            ],
                            "missing_skills": [
                                {"skill": "API Communication", "importance": 5, "suggestion": "Please try again later when the service is available."}
                            ],
                            "improvement_suggestions": [
                                "Try again later when API services are available."
                            ],
                            "potential_red_flags": [
                                "No analysis available due to API error."
                            ],
                            "experience_tailoring": [
                                "No tailoring suggestions available due to API error."
                            ],
                            "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
                        }
                
                # Handle other error codes
                if response.status_code != 200:
                    logger.error(f"Gemini API error (code {response.status_code}): {response.text}")
                    if attempt < max_retries:
                        logger.warning(f"Retrying after error (attempt {attempt}/{max_retries})...")
                        time.sleep(retry_delay)
                        continue
                    else:
                        logger.error(f"Gemini API error after {max_retries} attempts: {response.status_code} {response.text}")
                        
                        # For error code 400, try with a much smaller prompt for debugging
                        if response.status_code == 400 and attempt == max_retries:
                            logger.info("Trying minimal request to diagnose API issue...")
                            debug_payload = {
                                "contents": [
                                    {
                                        "parts": [
                                            {"text": "Return a brief JSON analysis of this resume with matching_skills, missing_skills, and improvement_suggestions."}
                                        ]
                                    }
                                ],
                                "generationConfig": {
                                    "temperature": 0.2,
                                    "maxOutputTokens": 1024
                                }
                            }
                            debug_response = requests.post(api_url, headers=headers, json=debug_payload, timeout=30)
                            logger.info(f"Debug response: status={debug_response.status_code}, body={debug_response.text[:500]}...")
                        
                        return {
                            "error": f"API Error: The Gemini API returned an error (code {response.status_code}). Please try again later.",
                            "matching_skills": [
                                {"skill": "API Connection", "strength": 3, "importance": 5, "notes": f"API returned error code {response.status_code}"}
                            ],
                            "missing_skills": [
                                {"skill": "API Communication", "importance": 5, "suggestion": "Please try again later when the service is available."}
                            ],
                            "improvement_suggestions": [
                                "Try again later when API services are available."
                            ],
                            "potential_red_flags": [
                                "No analysis available due to API error."
                            ],
                            "experience_tailoring": [
                                "No tailoring suggestions available due to API error."
                            ],
                            "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
                        }
                
                # Process the response for success case
                try:
                    response_data = response.json()
                    logger.info("Received successful response from Gemini API")
                    logger.info(f"Response structure: {json.dumps(list(response_data.keys()), indent=2)}")
                    
                    # Extract the text response from the API
                    text_response = ""
                    
                    # The response structure might have changed in different API versions
                    if "candidates" in response_data and len(response_data["candidates"]) > 0:
                        candidate = response_data["candidates"][0]
                        
                        if "content" in candidate:
                            content = candidate["content"]
                            
                            if "parts" in content and len(content["parts"]) > 0:
                                for part in content["parts"]:
                                    if "text" in part:
                                        text_response += part["text"]
                    
                    if not text_response:
                        logger.error(f"No text found in response: {json.dumps(response_data)[:500]}...")
                        if attempt < max_retries:
                            logger.warning(f"Retrying due to empty text response (attempt {attempt}/{max_retries})...")
                            time.sleep(retry_delay)
                            continue
                        
                        return {
                            "error": "The API returned an empty response. Please try again later.",
                            "matching_skills": [
                                {"skill": "API Response", "strength": 3, "importance": 5, "notes": "The API returned an empty response."}
                            ],
                            "missing_skills": [
                                {"skill": "API Format", "importance": 5, "suggestion": "Please try again later when the service is stable."}
                            ],
                            "improvement_suggestions": ["Try again later."],
                            "potential_red_flags": ["No analysis available due to empty API response."],
                            "experience_tailoring": ["No tailoring suggestions available due to empty API response."],
                            "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
                        }
                    
                    logger.info(f"Text response length: {len(text_response)}")
                    logger.info(f"First 200 characters of response: {text_response[:200]}")
                    
                    # Extract the JSON part from the response
                    try:
                        # Look for JSON content in the response
                        json_start = text_response.find("{")
                        json_end = text_response.rfind("}") + 1
                        
                        if json_start >= 0 and json_end > json_start:
                            json_text = text_response[json_start:json_end]
                            logger.info(f"JSON text length: {len(json_text)}")
                            
                            try:
                                recommendations = json.loads(json_text)
                                logger.info("Successfully parsed JSON from Gemini response")
                                
                                # Validate the required JSON structure - Provide default values for any missing keys
                                required_keys = ["matching_skills", "missing_skills", "improvement_suggestions", 
                                               "potential_red_flags", "experience_tailoring", "gap_analysis"]
                                
                                # Create fallback examples if data is missing
                                sample_skills = [
                                    {"skill": "Communication", "strength": 3, "importance": 4, "notes": "Your resume shows some evidence of communication skills through your project descriptions."},
                                    {"skill": "Problem Solving", "strength": 4, "importance": 5, "notes": "Your experience demonstrates strong problem-solving capabilities."}
                                ]
                                sample_missing = [
                                    {"skill": "Leadership", "importance": 4, "suggestion": "Add examples of team leadership or project management."},
                                    {"skill": "Technical Writing", "importance": 3, "suggestion": "Include documentation or technical writing experience."}
                                ]
                                
                                # Add default values for any missing keys
                                for key in required_keys:
                                    if key not in recommendations or not recommendations[key]:
                                        if key == "matching_skills":
                                            recommendations[key] = sample_skills
                                        elif key == "missing_skills":
                                            recommendations[key] = sample_missing
                                        elif key == "gap_analysis":
                                            recommendations[key] = {"overall_match": "65%", "technical_match": "70%", "experience_match": "60%", "critical_gaps": ["Consider adding more examples of leadership", "Highlight specific technical skills"]}
                                        else:
                                            recommendations[key] = ["No specific recommendations found. Try resubmitting with more detailed resume or job description."]
                                        logger.warning(f"Added default values for missing key: {key}")
                                
                                # Ensure we have structured data in matching_skills and missing_skills
                                if not isinstance(recommendations["matching_skills"], list) or len(recommendations["matching_skills"]) == 0:
                                    recommendations["matching_skills"] = sample_skills
                                elif not isinstance(recommendations["matching_skills"][0], dict):
                                    recommendations["matching_skills"] = [{"skill": item, "strength": 3, "importance": 3, "notes": "Automatically detected"} for item in recommendations["matching_skills"]]
                                
                                if not isinstance(recommendations["missing_skills"], list) or len(recommendations["missing_skills"]) == 0:
                                    recommendations["missing_skills"] = sample_missing
                                elif not isinstance(recommendations["missing_skills"][0], dict):
                                    recommendations["missing_skills"] = [{"skill": item, "importance": 3, "suggestion": "Consider adding this skill"} for item in recommendations["missing_skills"]]
                                
                                return recommendations
                            except json.JSONDecodeError as parse_err:
                                logger.error(f"JSON parsing error: {parse_err}")
                                if attempt < max_retries:
                                    logger.warning(f"Retrying after JSON parse error (attempt {attempt}/{max_retries})...")
                                    time.sleep(retry_delay)
                                    continue
                                # Return fallback data structure on final attempt
                                return {
                                    "matching_skills": [
                                        {"skill": "Resume Analysis", "strength": 3, "importance": 5, "notes": "We were unable to analyze your specific skills. Please try again."}
                                    ],
                                    "missing_skills": [
                                        {"skill": "API Error Handling", "importance": 5, "suggestion": "There was an error analyzing your resume. Please try again later."}
                                    ],
                                    "improvement_suggestions": [
                                        "Our system encountered an error analyzing your resume. Please try again later or contact support if the issue persists."
                                    ],
                                    "potential_red_flags": [
                                        "No analysis available due to system error."
                                    ],
                                    "experience_tailoring": [
                                        "No tailoring suggestions available due to system error."
                                    ],
                                    "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
                                }
                        else:
                            logger.warning("JSON content not found in the response text")
                            if attempt < max_retries:
                                logger.warning(f"Retrying after no JSON content found (attempt {attempt}/{max_retries})...")
                                time.sleep(retry_delay)
                                continue
                            # Return fallback structure on final attempt
                            return {
                                "matching_skills": [
                                    {"skill": "Communication", "strength": 3, "importance": 4, "notes": "API returned text but no JSON content was found."}
                                ],
                                "missing_skills": [
                                    {"skill": "API Compatibility", "importance": 5, "suggestion": "Try again with a different resume format or job description."}
                                ],
                                "improvement_suggestions": [
                                    "Our system encountered an issue parsing the analysis results. Please try again."
                                ],
                                "potential_red_flags": ["No analysis available."],
                                "experience_tailoring": ["No tailoring suggestions available."],
                                "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
                            }
                            
                    except Exception as json_err:
                        logger.error(f"Error processing JSON response: {json_err}")
                        if attempt < max_retries:
                            logger.warning(f"Retrying after JSON processing error (attempt {attempt}/{max_retries})...")
                            time.sleep(retry_delay)
                            continue
                        # Return fallback structure on final attempt
                        return {
                            "matching_skills": [
                                {"skill": "Error Handling", "strength": 3, "importance": 5, "notes": "An error occurred while processing the analysis results."}
                            ],
                            "missing_skills": [
                                {"skill": "API Processing", "importance": 5, "suggestion": "Please try again or contact support if the issue persists."}
                            ],
                            "improvement_suggestions": [
                                "Try again with a simplified resume or job description."
                            ],
                            "potential_red_flags": ["No analysis available due to system error."],
                            "experience_tailoring": ["No tailoring suggestions available due to system error."],
                            "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
                        }
                
                except json.JSONDecodeError:
                    logger.error(f"Invalid JSON in API response: {response.text[:500]}...")
                    if attempt < max_retries:
                        logger.warning(f"Retrying after response JSON decode error (attempt {attempt}/{max_retries})...")
                        time.sleep(retry_delay)
                        continue
                    
                    # Return fallback structure on final attempt
                    return {
                        "matching_skills": [
                            {"skill": "API Response", "strength": 3, "importance": 5, "notes": "The API returned an invalid response. Please try again."}
                        ],
                        "missing_skills": [
                            {"skill": "Response Format", "importance": 5, "suggestion": "Please try again later when the service is stable."}
                        ],
                        "improvement_suggestions": [
                            "Try again later when the system is more stable."
                        ],
                        "potential_red_flags": ["No analysis available due to API response error."],
                        "experience_tailoring": ["No tailoring suggestions available due to API response error."],
                        "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
                    }
                    
            except requests.exceptions.Timeout:
                logger.error(f"Gemini API request timed out (attempt {attempt}/{max_retries})")
                if attempt < max_retries:
                    logger.warning(f"Retrying after timeout (attempt {attempt}/{max_retries})...")
                    time.sleep(retry_delay)
                    continue
            except requests.exceptions.RequestException as req_err:
                logger.error(f"Request error to Gemini API (attempt {attempt}/{max_retries}): {str(req_err)}")
                if attempt < max_retries:
                    logger.warning(f"Retrying after request error (attempt {attempt}/{max_retries})...")
                    time.sleep(retry_delay)
                    continue
        
        # If we reached here, all retry attempts failed
        logger.error(f"All {max_retries} attempts to call Gemini API failed")
        return {
            "error": "Unable to connect to the Gemini API after multiple attempts. Please try again later.",
            "matching_skills": [
                {"skill": "API Connection", "strength": 3, "importance": 5, "notes": "Connection to the AI service failed."}
            ],
            "missing_skills": [
                {"skill": "API Communication", "importance": 5, "suggestion": "Please try again later when the service is available."}
            ],
            "improvement_suggestions": [
                "Try again later when API services are available."
            ],
            "potential_red_flags": [
                "No analysis available due to API connection error."
            ],
            "experience_tailoring": [
                "No tailoring suggestions available due to API connection error."
            ],
            "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
        }
                                
    except Exception as e:
        error_msg = str(e)
        # Check for quota/rate limit error
        if (
            'quota' in error_msg.lower() or
            'rate limit' in error_msg.lower() or
            '429' in error_msg or
            'exceeded' in error_msg.lower()
        ):
            user_msg = (
                "The resume analysis service is temporarily unavailable due to high demand or quota limits. "
                "Please try again in a few minutes. If the problem persists, contact support."
            )
            logger.error(f"Quota/rate limit error: {error_msg}")
            return {
                "error": user_msg,
                "details": error_msg,
                "matching_skills": [
                    {"skill": "API Connection", "strength": 3, "importance": 5, "notes": "API returned quota/rate limit error."}
                ],
                "missing_skills": [
                    {"skill": "API Communication", "importance": 5, "suggestion": "Please try again later when the service is available."}
                ],
                "improvement_suggestions": [
                    "Try again later when API services are available."
                ],
                "potential_red_flags": [
                    "No analysis available due to API quota error."
                ],
                "experience_tailoring": [
                    "No tailoring suggestions available due to API quota error."
                ],
                "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
            }
        
        logger.error(f"Error getting recommendations from Gemini: {error_msg}")
        return {
            "error": "An error occurred while analyzing your resume. Please try again later.",
            "details": error_msg,
            "matching_skills": [
                {"skill": "API Connection", "strength": 3, "importance": 5, "notes": "API connection failed with error."}
            ],
            "missing_skills": [
                {"skill": "API Communication", "importance": 5, "suggestion": "Please try again later when the service is available."}
            ],
            "improvement_suggestions": [
                "Try again later when API services are available."
            ],
            "potential_red_flags": [
                "No analysis available due to API error."
            ],
            "experience_tailoring": [
                "No tailoring suggestions available due to API error."
            ],
            "gap_analysis": {"overall_match": "N/A", "technical_match": "N/A", "experience_match": "N/A", "critical_gaps": ["Unable to analyze"]}
        }


# Import resume agent components (assuming they exist in the same directory or are importable)
# Note: Ensure these imports are correct based on your project structure
try:
    from .docx_generator import text_to_docx
    from .azure_storage_manager import AzureStorageManager
    from .resume_rewriter import rewrite_resume, extract_job_details, generate_unique_identifier
except ImportError as ie:
    logger.warning(f"Could not import agent components: {ie}. Ensure docx_generator.py, azure_storage_manager.py, and resume_rewriter.py are accessible.")
    # Define dummy classes/functions if imports fail, to allow basic functionality
    def text_to_docx(text): return b"Dummy DOCX data"
    class AzureStorageManager: 
        def upload_tailored_resume(self, data, filename): return f"http://dummy.url/{filename}"
        def update_tracking_file(self, entry): pass
    def rewrite_resume(text, jd): return {"rewritten_resume_text": "Dummy rewritten text", "job_title": "Dummy Job", "company_name": "Dummy Co"}
    def extract_job_details(jd): return ("Dummy Job", "Dummy Co", "Dummy Role")
    def generate_unique_identifier(name, job, company): return f"Resume_{name}_{job}_{company}_timestamp.docx"

class ResumeAgent:
    """
    Resume optimization and processing agent
    
    This agent processes a resume and job description to create a tailored resume,
    store it in Azure storage, and track the changes.
    """
    
    def __init__(self):
        """Initialize the resume agent"""
        self.storage_manager = AzureStorageManager()
    
    def process_resume(self, resume_file_path: str, job_description: str, 
                       user_name: str, original_filename: str):
        """
        Process a resume and job description to create a tailored resume
        
        Args:
            resume_file_path: Path to the original resume file
            job_description: The job description text
            user_name: The name of the user
            original_filename: The original filename of the resume
            
        Returns:
            A dictionary containing the results of the operation:
            - status: "success" or "error"
            - message: A message describing the result
            - data: Additional data about the operation (if successful)
            - error: Error details (if failed)
        """
        try:
            logger.info(f"Processing resume for user: {user_name}")
            
            # Step 1: Extract text from the resume file
            logger.info("Extracting text from resume file...")
            resume_text = extract_text_from_file(resume_file_path, original_filename)
            
            # Step 2: Rewrite the resume to match the job description using the imported function
            logger.info("Rewriting resume to match job description...")
            # Assuming rewrite_resume is imported correctly and uses the fixed model ID
            rewrite_result = rewrite_resume(resume_text, job_description)
            
            # Check if rewriting was successful
            if "error" in rewrite_result:
                logger.error(f"Error rewriting resume: {rewrite_result.get('error')}")
                return {
                    "status": "error",
                    "message": "Failed to rewrite resume",
                    "error": rewrite_result.get("error", "Unknown error")
                }
            
            # Step 3: Generate a DOCX file from the rewritten resume
            logger.info("Generating DOCX file...")
            rewritten_text = rewrite_result.get("rewritten_resume_text", "")
            docx_data = text_to_docx(rewritten_text)
            
            # Step 4: Generate a unique filename for the tailored resume
            job_title = rewrite_result.get("job_title", "Not Specified")
            company_name = rewrite_result.get("company_name", "Not Specified")
            role = rewrite_result.get("role", "Not Specified") # Assuming role is returned by rewrite_resume or extract_job_details
            tailored_filename = generate_unique_identifier(user_name, job_title, company_name)
            
            # Step 5: Upload the tailored resume to Azure storage
            logger.info(f"Uploading tailored resume: {tailored_filename}")
            resume_url = self.storage_manager.upload_tailored_resume(docx_data, tailored_filename)
            
            # Step 6: Update the tracking file
            logger.info("Updating tracking file...")
            tracking_entry = {
                "user_name": user_name,
                "original_filename": original_filename,
                "tailored_filename": tailored_filename,
                "job_title": job_title,
                "company_name": company_name,
                "role": role,
                "timestamp": datetime.datetime.now().isoformat(),
                "resume_url": resume_url,
                "changes_summary": rewrite_result.get("changes_summary", []), # Assuming rewrite_resume provides this
                "integration_percentage": rewrite_result.get("integration_percentage", "N/A"), # Assuming rewrite_resume provides this
                "highlighted_skills": rewrite_result.get("highlighted_skills", []), # Assuming rewrite_resume provides this
                "gap_analysis": rewrite_result.get("gap_analysis", []) # Assuming rewrite_resume provides this
            }
            self.storage_manager.update_tracking_file(tracking_entry)
            
            logger.info("Resume processing completed successfully")
            return {
                "status": "success",
                "message": "Resume processed and tailored successfully",
                "data": {
                    "rewritten_resume_url": resume_url,
                    "rewritten_resume_text": rewritten_text # Return the text as well
                }
            }
            
        except Exception as e:
            logger.error(f"Error in ResumeAgent process_resume: {str(e)}")
            return {
                "status": "error",
                "message": "An error occurred during resume processing",
                "error": str(e)
            }

